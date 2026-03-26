"""
parsers/docx_parser.py  –  DOCX text extractor
"""
from typing import Dict, Any
from utils.logger import get_logger

log = get_logger("docx_parser")


def extract_docx(file_bytes: bytes) -> Dict[str, Any]:
    """Extract text from a DOCX file."""
    try:
        import docx
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    import io
    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full_text  = "\n".join(paragraphs)

    # Tables
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                table_texts.append(" | ".join(row_texts))
    if table_texts:
        full_text += "\n" + "\n".join(table_texts)

    log.info("DOCX extracted: paragraphs=%d chars=%d", len(paragraphs), len(full_text))

    return {
        "full_text":     full_text,
        "left_text":     full_text,
        "right_text":    "",
        "pages":         1,
        "is_two_column": False,
        "raw_blocks":    [],
    }
